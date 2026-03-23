import asyncio
import aiohttp
import os
import io
import re
import tempfile
from pathlib import Path
from fastapi import UploadFile, HTTPException
from docx import Document
import pdfplumber
import docx2txt
from settings import settings, logger
from typing import List


# ── PDF OCR 标题检测（参考 PDFconvert）────────────────────────

# 中文句子结尾标点，出现则不视为标题
_SENTENCE_ENDINGS = re.compile(r'[。，；：、！？,.;:!?…]$')

# 典型中文标题前缀：第X章/节、一、（一）、1. 等
_HEADING_PREFIX = re.compile(
    r'^(第[一二三四五六七八九十百]+[章节条款]|[一二三四五六七八九十]+[、．.]|\d+[.、）)]\s*|（[一二三四五六七八九十]+）)'
)


def _is_heading(line: str) -> bool:
    """判断一行文本是否为标题（规则来自 PDFconvert）。"""
    if _HEADING_PREFIX.match(line):
        return True
    if len(line) < 25 and not _SENTENCE_ENDINGS.search(line):
        return True
    return False


def _text_to_markdown(text: str) -> str:
    """将 OCR 原始文本转换为带 ## 标题标记的 Markdown 字符串。"""
    lines = text.split("\n")
    md = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if _is_heading(line):
            md.append("## " + line)
        else:
            md.append(line)
    return "\n\n".join(md)


def _pdf_to_markdown_sync(pdf_path: Path) -> str:
    """
    同步 PDF → Markdown 转换。
    优先用 pdfplumber 直接提取文字（速度快、省内存）；
    若提取内容不足（扫描版 PDF），降级为逐页 OCR（150 DPI）。
    在线程池中调用，避免阻塞事件循环。
    """
    # ── 尝试 pdfplumber 直接提取文字 ──────────────────────────────
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            pages_text = [p.extract_text() or "" for p in pdf.pages]
        full_text = "\n".join(pages_text).replace('\x00', '').strip()
        avg_chars = len(full_text) / max(len(pages_text), 1)
        if avg_chars >= 50:   # 每页至少 50 字符，认为是文字型 PDF
            logger.info("PDF 文字提取成功（均 %.0f 字/页）: %s", avg_chars, pdf_path.name)
            return _text_to_markdown(full_text)
        logger.info("PDF 文字提取内容不足（均 %.0f 字/页），改用 OCR: %s", avg_chars, pdf_path.name)
    except Exception as e:
        logger.warning("pdfplumber 提取失败，改用 OCR: %s — %s", pdf_path.name, e)

    # ── 降级：逐页 OCR，150 DPI，避免全量加载占用过多内存 ────────
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        raise RuntimeError(
            f"PDF OCR 依赖未安装: {e}。请运行: pip install pdf2image pytesseract"
        ) from e

    with pdfplumber.open(str(pdf_path)) as pdf:
        page_count = len(pdf.pages)

    text = ""
    for i in range(1, page_count + 1):
        images = convert_from_path(str(pdf_path), dpi=150, first_page=i, last_page=i)
        for img in images:
            text += pytesseract.image_to_string(img, lang="chi_sim+eng") + "\n"
        del images  # 释放当页图片内存
    return _text_to_markdown(text)


async def pdf_to_markdown(pdf_path: Path) -> str:
    """异步包装：在线程池中执行 PDF OCR 转换，返回 Markdown 文本。"""
    return await asyncio.to_thread(_pdf_to_markdown_sync, pdf_path)


def _epub_to_markdown_sync(epub_path: Path) -> str:
    """
    同步 epub → Markdown 转换：提取各章节 HTML，将 h1-h6 转为 ## 标题。
    在线程池中调用，避免阻塞事件循环。
    """
    try:
        import ebooklib
        from ebooklib import epub as epublib
    except ImportError as e:
        raise RuntimeError(
            f"epub 依赖未安装: {e}。请运行: pip install ebooklib"
        ) from e

    from html.parser import HTMLParser

    class _EpubHTMLParser(HTMLParser):
        _HEADING = {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        _BLOCK   = {'p', 'div', 'li', 'td', 'th'}

        def __init__(self):
            super().__init__()
            self._buf = []
            self._in_heading = False

        def handle_starttag(self, tag, attrs):
            tag = tag.lower()
            if tag in self._HEADING:
                self._buf.append('\n## ')
                self._in_heading = True
            elif tag == 'br':
                self._buf.append('\n')
            elif tag in self._BLOCK:
                self._buf.append('\n')

        def handle_endtag(self, tag):
            tag = tag.lower()
            if tag in self._HEADING:
                self._buf.append('\n')
                self._in_heading = False
            elif tag in self._BLOCK:
                self._buf.append('\n')

        def handle_data(self, data):
            self._buf.append(data)

        def result(self):
            return ''.join(self._buf)

    book = epublib.read_epub(str(epub_path), options={'ignore_ncx': True})
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        raw_html = item.get_content().decode('utf-8', errors='ignore')
        parser = _EpubHTMLParser()
        parser.feed(raw_html)
        parts.append(parser.result())

    md = '\n\n'.join(parts)
    md = re.sub(r'\n{3,}', '\n\n', md)
    return md.strip()


async def epub_to_markdown(epub_path: Path) -> str:
    """异步包装：在线程池中执行 epub → Markdown 转换。"""
    return await asyncio.to_thread(_epub_to_markdown_sync, epub_path)


def split_markdown_chunks(md_text: str, max_size: int = 800) -> List[str]:
    """
    按 ## 标题边界拆分 Markdown 为语义 chunks。
    超长 section 再按段落细分，确保每块不超过 max_size 字符。
    """
    sections = re.split(r'(?=^## )', md_text, flags=re.MULTILINE)
    chunks = []
    for section in sections:
        section = section.strip()
        if not section:
            continue
        if len(section) <= max_size:
            chunks.append(section)
        else:
            sub_paragraphs = split_into_paragraphs(section)
            chunks.extend(group_paragraphs(sub_paragraphs, max_size=max_size))
    return chunks


# 使用 Google Search 获取最新网络信息
async def fetch_from_web(query: str):
    search_url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": settings.google_api_key,
        "cx": settings.google_cx,
        "q": query,
        "num": 5  # 获取前5个结果
    }
    proxy = os.environ.get("HTTP_PROXY") or os.environ.get("http_proxy")
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(search_url, params=params, proxy=proxy,
                                   timeout=aiohttp.ClientTimeout(total=10)) as response:
                if response.status == 200:
                    data = await response.json()
                    results = []
                    for item in data.get("items", []):
                        title = item.get("title")
                        snippet = item.get("snippet")
                        link = item.get("link")
                        results.append(f"Title: {title}\nSnippet: {snippet}\nLink: {link}")
                    return "\n\n".join(results)
                else:
                    return ""
    except Exception as e:
        logger.warning("Web 搜索失败: %s", e)
        return ""


# ── 文档解析 ──────────────────────────────────────────────────

def _parse_document_sync(path: Path) -> str:
    """同步读取并解析文档，在线程池中执行以避免阻塞事件循环。"""
    data = path.read_bytes()
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    elif suffix == ".docx":
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    elif suffix == ".doc":
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        text = docx2txt.process(tmp_path)
        os.remove(tmp_path)
        return text
    elif suffix == ".pdf":
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)
    elif suffix == ".epub":
        # epub 走 Markdown 路径，此处提供纯文本兜底
        return _epub_to_markdown_sync(path)
    else:
        raise HTTPException(status_code=400, detail="暂不支持该文件格式")


async def parse_document(path: Path) -> str:
    """异步包装器：在线程池中执行同步文档解析。"""
    return await asyncio.to_thread(_parse_document_sync, path)


# ── 段落感知分块 ──────────────────────────────────────────────

def split_into_paragraphs(text: str) -> List[str]:
    """按双换行拆段，兼容中英文。"""
    parts = re.split(r'\n\s*\n', text)
    return [p.strip() for p in parts if p.strip()]


def group_paragraphs(paragraphs: List[str], max_size: int = 800) -> List[str]:
    """将短段落合并为不超过 max_size 字符的块；超大段落直接单独成块。"""
    chunks, current, current_len = [], [], 0
    for p in paragraphs:
        if current_len + len(p) > max_size and current:
            chunks.append("\n\n".join(current))
            current, current_len = [], 0
        current.append(p)
        current_len += len(p)
    if current:
        chunks.append("\n\n".join(current))
    return chunks


# ── Gemini 上下文增强 ─────────────────────────────────────────

def enrich_chunks_with_context(doc_text: str, filename: str, chunks: List[str]) -> List[str]:
    """
    为每个 chunk 拼接上下文头，使 embedding 携带来源和位置信息。
    使用文档开头的自然文字作为上下文摘要，无需 Gemini API 调用。
    """
    # 取文档前 300 字作为上下文提示（去除多余空白）
    doc_intro = " ".join(doc_text[:500].split())[:300]
    total = len(chunks)
    enriched = []
    for i, chunk in enumerate(chunks):
        header = (
            f"[来源文件：{filename}。文档开头：{doc_intro}。"
            f"位置：第{i + 1}段，共{total}段。]\n\n"
        )
        enriched.append(header + chunk)
    return enriched


# ── 旧版接口（保留兼容性）────────────────────────────────────

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


def parse_text_from_bytes(data: bytes, suffix: str, chunk_size: int = 1000) -> List[str]:
    suffix = suffix.lower()
    if suffix == ".txt":
        text = data.decode("utf-8", errors="ignore")
    elif suffix == ".docx":
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
    elif suffix == ".doc":
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        text = docx2txt.process(tmp_path)
        os.remove(tmp_path)
    elif suffix == ".pdf":
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = "\n".join(pages)
    else:
        raise HTTPException(status_code=400, detail="暂不支持该文件格式")
    return chunk_text(text, chunk_size=chunk_size)


async def extract_text_chunks(file: UploadFile) -> List[str]:
    suffix = os.path.splitext(file.filename)[1].lower()
    data = await file.read()
    return parse_text_from_bytes(data, suffix)


async def extract_text_chunks_path(path: Path) -> List[str]:
    suffix = path.suffix.lower()
    data = path.read_bytes()
    return parse_text_from_bytes(data, suffix)
